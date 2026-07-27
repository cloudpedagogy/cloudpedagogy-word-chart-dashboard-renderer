import tempfile
import unittest
from pathlib import Path
from docx import Document
import sys
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from render_chart_dashboard import InputError, parse_document, render

ROOT = Path(__file__).resolve().parents[1]
SAMPLE = ROOT / "examples" / "chart_dashboard_example.docx"

def make_doc(charts, data):
    d=Document()
    d.add_heading("CHARTS",1); t=d.add_table(rows=1,cols=3)
    for c,v in zip(t.rows[0].cells,["chart_id","title","chart_type"]): c.text=v
    for row in charts:
        for c,v in zip(t.add_row().cells,row): c.text=str(v)
    d.add_heading("DATA",1); t=d.add_table(rows=1,cols=4)
    for c,v in zip(t.rows[0].cells,["chart_id","series","x","y"]): c.text=v
    for row in data:
        for c,v in zip(t.add_row().cells,row): c.text=str(v)
    p=Path(tempfile.mkdtemp())/"input.docx";d.save(p);return p

class Tests(unittest.TestCase):
    def test_sample_parses(self):
        settings,charts,data,qa=parse_document(SAMPLE)
        self.assertEqual(len(charts),6);self.assertGreater(len(data),30);self.assertFalse(qa.errors)
    def test_all_types(self):
        _,charts,_,qa=parse_document(SAMPLE)
        self.assertEqual({c["chart_type"] for c in charts},{"bar","line","scatter","area","pie","donut"});self.assertFalse(qa.errors)
    def test_alias_headers(self):
        p=make_doc([("one","Alias","bar")],[("one","A","Alpha",3)])
        _,charts,data,qa=parse_document(p);self.assertEqual(charts[0]["chart_id"],"one");self.assertEqual(data[0]["y"],3);self.assertFalse(qa.errors)
    def test_duplicate_chart(self):
        p=make_doc([("one","A","bar"),("one","B","line")],[("one","A","X",2)])
        *_,qa=parse_document(p);self.assertTrue(any("duplicate" in e for e in qa.errors))
    def test_bad_number(self):
        p=make_doc([("one","A","bar")],[("one","A","X","not a number")])
        *_,qa=parse_document(p);self.assertTrue(any("valid number" in e for e in qa.errors))
    def test_unknown_reference(self):
        p=make_doc([("one","A","bar")],[("missing","A","X",2)])
        *_,qa=parse_document(p);self.assertTrue(any("does not match" in e for e in qa.errors))
    def test_offline_render(self):
        with tempfile.TemporaryDirectory() as td:
            out=Path(td)/"out";render(SAMPLE,out)
            html=(out/"index.html").read_text()
            self.assertIn("Plotly",html);self.assertNotIn('src="https://cdn.plot.ly',html);self.assertTrue((out/"data.json").exists())

if __name__=="__main__": unittest.main()
