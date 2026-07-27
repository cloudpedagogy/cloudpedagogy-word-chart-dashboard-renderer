#!/usr/bin/env python3
from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "examples" / "chart_dashboard_example.docx"

def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0]._element.get_or_add_trPr().append(__import__("docx").oxml.OxmlElement("w:tblHeader"))
    for i, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]; cell.width = Cm(width); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True; run.font.color.rgb = RGBColor(255,255,255); run.font.size = Pt(8)
        cell._tc.get_or_add_tcPr().append(__import__("docx").oxml.parse_xml('<w:shd {} w:fill="2457D6"/>'.format(__import__("docx").oxml.ns.nsdecls("w"))))
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].width = Cm(widths[i]); cells[i].text = str(value); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs: run.font.size = Pt(7.5)
        trPr = cells[0]._tc.getparent().get_or_add_trPr()
        cant = __import__("docx").oxml.OxmlElement("w:cantSplit"); trPr.append(cant)
    doc.add_paragraph()

doc = Document()
sec = doc.sections[0]; sec.orientation = WD_ORIENT.LANDSCAPE; sec.page_width, sec.page_height = sec.page_height, sec.page_width
sec.left_margin = sec.right_margin = Cm(1.3); sec.top_margin = sec.bottom_margin = Cm(1.3)
styles = doc.styles
styles["Normal"].font.name = "Arial"; styles["Normal"].font.size = Pt(9)
title = doc.add_heading("Interactive Chart Dashboard — Example Input", 0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("Edit the tables below, then run render_chart_dashboard.py. The DATA table uses a flexible long format: one row per plotted value.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("SETTINGS", 1)
add_table(doc, ["setting","value"], [
    ("title","Digital Learning Evidence Dashboard"),("subtitle","Illustrative quantitative data entered in Word"),
    ("theme","light"),("columns","2"),("show_data_table","yes"),("responsive","yes")
], [7,19])

doc.add_heading("CHARTS", 1)
charts = [
("access","Resource access by type","bar","Students accessing each resource","Resource type","Students","yes","yes","input","","420"),
("engagement","Monthly learning activity","line","Views across the academic year","Month","Views","no","yes","input","","420"),
("relationship","Engagement and assessment","scatter","Bubble size represents enrolment","Engagement rate (%)","Mean mark (%)","no","yes","input","","420"),
("completion","Completion trend","area","Cumulative completion by pathway","Week","Students completed","yes","yes","input","","420"),
("activity_mix","Activity mix","pie","Share of all course activities","","","no","yes","y descending","","420"),
("delivery","Delivery mode","donut","Student enrolment by mode","","","no","yes","input","0.58","420"),
]
add_table(doc, ["chart_id","title","chart_type","subtitle","x_label","y_label","stacked","show_legend","sort","donut_hole","height"], charts, [2.4,4.2,2.1,4.7,3.2,3.2,1.7,2,2.4,2.1,1.6])

doc.add_heading("DATA", 1)
data=[]
for series, vals, color in [("Accessed",[410,355,290,245],"#2457D6"),("Not accessed",[40,95,160,205],"#E45756")]:
    for x,y in zip(["Moodle Book","Quiz","Video","Forum"],vals): data.append(("access",series,x,y,"","",color))
for series,vals,color in [("2025–26",[820,1050,1320,1190,1480,1660],"#2457D6"),("2024–25",[710,910,1110,1030,1290,1410],"#2A9D8F")]:
    for x,y in zip(["Oct","Nov","Dec","Jan","Feb","Mar"],vals): data.append(("engagement",series,x,y,"","",color))
for label,x,y,size,color in [("Biostatistics",82,71,260,"#2457D6"),("Epidemiology",76,68,220,"#2A9D8F"),("Health Policy",69,66,180,"#F2A541"),("Research Methods",88,74,300,"#7B61A8"),("Economics",62,64,140,"#E45756")]:
    data.append(("relationship","Modules",x,y,size,label,color))
for series,vals,color in [("Distance learning",[12,38,75,118,165,203],"#2457D6"),("Intensive",[8,31,61,91,121,146],"#2A9D8F")]:
    for x,y in zip(["1","2","3","4","5","6"],vals): data.append(("completion",series,x,y,"","",color))
for x,y,color in [("Books",37,"#2457D6"),("Resources",42,"#2A9D8F"),("Quizzes",18,"#F2A541"),("Forums",14,"#7B61A8"),("Other",20,"#A0AEC0")]:
    data.append(("activity_mix","Activities",x,y,"","",color))
for x,y,color in [("Distance learning",420,"#2457D6"),("Intensive",180,"#2A9D8F"),("Blended",95,"#F2A541")]:
    data.append(("delivery","Students",x,y,"","",color))
add_table(doc, ["chart_id","series","x","y","size","label","color"], data, [2.8,3.6,4.3,2.2,2.2,5.2,2.8])
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
