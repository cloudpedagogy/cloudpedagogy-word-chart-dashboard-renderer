#!/usr/bin/env python3
"""Render a multi-chart Plotly dashboard from structured Word tables."""
from __future__ import annotations

import argparse
import json
import math
import re
import shutil
import sys
from dataclasses import dataclass, field
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Any

from docx import Document

VERSION = "1.0.0"
SUPPORTED_TYPES = {"bar", "line", "scatter", "area", "pie", "donut"}
TRUE_VALUES = {"1", "true", "yes", "y", "on"}
FALSE_VALUES = {"0", "false", "no", "n", "off"}

TABLE_ALIASES = {
    "settings": {"settings", "configuration", "config", "dashboard settings"},
    "charts": {"charts", "chart definitions", "visualisations", "visualizations"},
    "data": {"data", "chart data", "values", "records"},
}
HEADER_ALIASES = {
    "chart_id": {"chart id", "chart_id", "chart", "visualisation id", "visualization id"},
    "chart_type": {"chart type", "chart_type", "type", "visualisation type", "visualization type"},
    "title": {"title", "chart title", "name"},
    "subtitle": {"subtitle", "sub title", "description"},
    "x_label": {"x label", "x_label", "x axis label", "x-axis label"},
    "y_label": {"y label", "y_label", "y axis label", "y-axis label"},
    "stacked": {"stacked", "stack", "stacking"},
    "show_legend": {"show legend", "show_legend", "legend"},
    "sort": {"sort", "sort order", "order"},
    "donut_hole": {"donut hole", "donut_hole", "hole", "hole size"},
    "height": {"height", "chart height", "height px"},
    "x": {"x", "x value", "category", "date", "x_value"},
    "y": {"y", "y value", "value", "amount", "y_value"},
    "series": {"series", "group", "category series", "trace"},
    "size": {"size", "marker size", "bubble size"},
    "label": {"label", "point label", "hover label"},
    "color": {"color", "colour", "series color", "series colour"},
    "setting": {"setting", "key", "option"},
    "setting_value": {"value", "setting value", "setting_value"},
}


class InputError(ValueError):
    pass


@dataclass
class QA:
    errors: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    def error(self, message: str) -> None:
        self.errors.append(message)

    def warn(self, message: str) -> None:
        self.warnings.append(message)


def normalise(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", value.strip().lower()).strip()


def canonical_header(value: str) -> str:
    key = normalise(value)
    for canonical, aliases in HEADER_ALIASES.items():
        if key in {normalise(a) for a in aliases}:
            return canonical
    return key.replace(" ", "_")


def clean_cell(cell: Any) -> str:
    return "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip()).strip()


def find_heading_before(table: Any) -> str:
    element = table._element.getprevious()
    while element is not None:
        if element.tag.endswith("}p"):
            texts = [node.text or "" for node in element.iter() if node.tag.endswith("}t")]
            text = "".join(texts).strip()
            if text:
                return normalise(text)
        element = element.getprevious()
    return ""


def table_kind(table: Any) -> str | None:
    heading = find_heading_before(table)
    for kind, aliases in TABLE_ALIASES.items():
        if heading in {normalise(a) for a in aliases}:
            return kind
    if not table.rows:
        return None
    headers = {canonical_header(clean_cell(c)) for c in table.rows[0].cells}
    if {"setting", "setting_value"} <= headers:
        return "settings"
    if {"chart_id", "chart_type"} <= headers:
        return "charts"
    if {"chart_id", "x", "y"} <= headers:
        return "data"
    return None


def read_rows(table: Any, kind: str | None = None) -> list[dict[str, str]]:
    if not table.rows:
        return []
    raw_headers = [clean_cell(c) for c in table.rows[0].cells]
    headers = [canonical_header(v) for v in raw_headers]
    if kind == "settings":
        headers = ["setting_value" if normalise(v) in {"value", "setting value"} else h for v, h in zip(raw_headers, headers)]
    if len(headers) != len(set(headers)):
        raise InputError(f"Duplicate column headings found: {headers}")
    rows = []
    for row in table.rows[1:]:
        values = [clean_cell(c) for c in row.cells]
        if not any(values):
            continue
        rows.append(dict(zip(headers, values)))
    return rows


def parse_bool(value: str, default: bool, qa: QA, context: str) -> bool:
    if not value.strip():
        return default
    key = normalise(value)
    if key in TRUE_VALUES:
        return True
    if key in FALSE_VALUES:
        return False
    qa.warn(f"{context}: '{value}' is not a recognised boolean; using {default}.")
    return default


def parse_number(value: str, qa: QA, context: str, required: bool = False) -> float | None:
    text = value.strip().replace(",", "")
    if not text:
        if required:
            qa.error(f"{context}: a numeric value is required.")
        return None
    try:
        number = float(text)
        if not math.isfinite(number):
            raise ValueError
        return number
    except ValueError:
        qa.error(f"{context}: '{value}' is not a valid number.")
        return None


def parse_document(path: Path) -> tuple[dict[str, Any], list[dict[str, Any]], list[dict[str, Any]], QA]:
    qa = QA()
    doc = Document(path)
    grouped: dict[str, list[dict[str, str]]] = {"settings": [], "charts": [], "data": []}
    for index, table in enumerate(doc.tables, 1):
        kind = table_kind(table)
        if kind:
            try:
                grouped[kind].extend(read_rows(table, kind))
            except InputError as exc:
                qa.error(f"Table {index}: {exc}")
        else:
            qa.warn(f"Table {index} was ignored because its purpose could not be identified.")

    settings: dict[str, Any] = {
        "title": "Chart Dashboard",
        "subtitle": "",
        "theme": "light",
        "columns": 2,
        "show_data_table": True,
        "responsive": True,
    }
    for i, row in enumerate(grouped["settings"], 2):
        key = normalise(row.get("setting", "")).replace(" ", "_")
        value = row.get("setting_value", "")
        if not key:
            qa.warn(f"SETTINGS row {i}: blank setting ignored.")
            continue
        settings[key] = value
    try:
        settings["columns"] = max(1, min(3, int(settings.get("columns", 2))))
    except (ValueError, TypeError):
        qa.warn("SETTINGS: columns must be 1, 2 or 3; using 2.")
        settings["columns"] = 2
    settings["show_data_table"] = parse_bool(str(settings.get("show_data_table", "")), True, qa, "SETTINGS show_data_table")
    settings["responsive"] = parse_bool(str(settings.get("responsive", "")), True, qa, "SETTINGS responsive")
    if normalise(str(settings.get("theme", "light"))) not in {"light", "dark"}:
        qa.warn("SETTINGS: theme must be light or dark; using light.")
        settings["theme"] = "light"
    else:
        settings["theme"] = normalise(str(settings["theme"]))

    charts: list[dict[str, Any]] = []
    ids: set[str] = set()
    for i, row in enumerate(grouped["charts"], 2):
        cid = row.get("chart_id", "").strip()
        ctype = normalise(row.get("chart_type", ""))
        if not cid:
            qa.error(f"CHARTS row {i}: chart_id is required.")
            continue
        if cid in ids:
            qa.error(f"CHARTS row {i}: duplicate chart_id '{cid}'.")
            continue
        ids.add(cid)
        if ctype not in SUPPORTED_TYPES:
            qa.error(f"CHARTS row {i}: unsupported chart_type '{row.get('chart_type', '')}'.")
            continue
        hole_default = 0.5 if ctype == "donut" else 0.0
        hole = parse_number(row.get("donut_hole", ""), qa, f"CHARTS row {i} donut_hole") if row.get("donut_hole", "").strip() else hole_default
        if hole is not None and not 0 <= hole <= 0.9:
            qa.warn(f"CHARTS row {i}: donut_hole must be 0–0.9; using {hole_default}.")
            hole = hole_default
        try:
            height = int(row.get("height", "") or 420)
            if height < 260 or height > 1000:
                raise ValueError
        except ValueError:
            qa.warn(f"CHARTS row {i}: height must be 260–1000; using 420.")
            height = 420
        sort = normalise(row.get("sort", "") or "input")
        if sort not in {"input", "x ascending", "x descending", "y ascending", "y descending"}:
            qa.warn(f"CHARTS row {i}: unknown sort '{row.get('sort', '')}'; using input.")
            sort = "input"
        charts.append({
            "chart_id": cid,
            "title": row.get("title", "").strip() or cid,
            "subtitle": row.get("subtitle", "").strip(),
            "chart_type": ctype,
            "x_label": row.get("x_label", "").strip(),
            "y_label": row.get("y_label", "").strip(),
            "stacked": parse_bool(row.get("stacked", ""), False, qa, f"CHARTS row {i} stacked"),
            "show_legend": parse_bool(row.get("show_legend", ""), True, qa, f"CHARTS row {i} show_legend"),
            "sort": sort,
            "donut_hole": hole,
            "height": height,
        })

    if not charts:
        qa.error("No valid CHARTS rows were found.")

    data: list[dict[str, Any]] = []
    counts = {cid: 0 for cid in ids}
    for i, row in enumerate(grouped["data"], 2):
        cid = row.get("chart_id", "").strip()
        if cid not in ids:
            qa.error(f"DATA row {i}: chart_id '{cid}' does not match a CHARTS row.")
            continue
        chart = next((c for c in charts if c["chart_id"] == cid), None)
        if not chart:
            continue
        x = row.get("x", "").strip()
        y = parse_number(row.get("y", ""), qa, f"DATA row {i} y", required=True)
        if not x:
            qa.error(f"DATA row {i}: x is required.")
        if y is None or not x:
            continue
        size = parse_number(row.get("size", ""), qa, f"DATA row {i} size") if row.get("size", "").strip() else None
        if size is not None and size <= 0:
            qa.warn(f"DATA row {i}: size must be positive; it will be ignored.")
            size = None
        data.append({
            "chart_id": cid,
            "series": row.get("series", "").strip() or "Value",
            "x": x,
            "y": y,
            "size": size,
            "label": row.get("label", "").strip(),
            "color": row.get("color", "").strip(),
            "row": i,
        })
        counts[cid] += 1

    for cid, count in counts.items():
        if count == 0 and any(c["chart_id"] == cid for c in charts):
            qa.warn(f"Chart '{cid}' has no valid DATA rows.")
    return settings, charts, data, qa


def payload(settings: dict[str, Any], charts: list[dict[str, Any]], data: list[dict[str, Any]]) -> dict[str, Any]:
    return {"schema_version": "1.0", "generator_version": VERSION, "settings": settings, "charts": charts, "data": data}


def build_html(bundle: dict[str, Any], plotly_js: str) -> str:
    encoded = json.dumps(bundle, ensure_ascii=False).replace("</", "<\\/")
    title = escape(str(bundle["settings"].get("title", "Chart Dashboard")))
    theme = bundle["settings"].get("theme", "light")
    return f"""<!doctype html>
<html lang="en" data-theme="{theme}">
<head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{title}</title>
<style>
:root{{--bg:#f4f7fb;--card:#fff;--text:#172033;--muted:#637083;--line:#dce3ee;--accent:#2457d6;--shadow:0 8px 24px #18345a14}}
[data-theme="dark"]{{--bg:#111827;--card:#1f2937;--text:#f3f4f6;--muted:#b7c0ce;--line:#374151;--accent:#7da6ff;--shadow:0 8px 24px #0005}}
*{{box-sizing:border-box}}body{{margin:0;background:var(--bg);color:var(--text);font:15px/1.5 system-ui,-apple-system,Segoe UI,sans-serif}}
header{{padding:28px clamp(18px,4vw,56px) 18px;background:var(--card);border-bottom:1px solid var(--line)}}h1{{margin:0;font-size:clamp(1.7rem,4vw,2.5rem)}}header p{{margin:.4rem 0 0;color:var(--muted)}}
.toolbar{{display:flex;gap:10px;flex-wrap:wrap;padding:16px clamp(18px,4vw,56px)}}button,input{{font:inherit;border:1px solid var(--line);border-radius:8px;padding:9px 12px;background:var(--card);color:var(--text)}}input{{min-width:230px}}button{{cursor:pointer}}button:hover{{border-color:var(--accent)}}
main{{display:grid;grid-template-columns:repeat(var(--cols),minmax(0,1fr));gap:18px;padding:0 clamp(18px,4vw,56px) 40px}}.card{{min-width:0;background:var(--card);border:1px solid var(--line);border-radius:14px;box-shadow:var(--shadow);overflow:hidden}}.card-head{{padding:18px 20px 0}}h2{{margin:0;font-size:1.15rem}}.subtitle{{margin:3px 0 0;color:var(--muted);font-size:.9rem}}.plot{{width:100%}}
details{{border-top:1px solid var(--line);padding:10px 18px 16px}}summary{{cursor:pointer;color:var(--accent)}}table{{width:100%;border-collapse:collapse;margin-top:10px;font-size:.85rem}}th,td{{text-align:left;padding:7px;border-bottom:1px solid var(--line)}}.empty{{padding:40px;text-align:center;color:var(--muted)}}
@media(max-width:900px){{main{{grid-template-columns:1fr}}}}@media print{{.toolbar{{display:none}}main{{display:block}}.card{{break-inside:avoid;margin-bottom:16px;box-shadow:none}}}}
</style>
<script>{plotly_js}</script>
</head>
<body>
<header><h1 id="dashboard-title"></h1><p id="dashboard-subtitle"></p></header>
<div class="toolbar"><input id="search" type="search" placeholder="Filter charts…" aria-label="Filter charts"><button id="theme">Switch theme</button><button id="download">Download data</button><button onclick="window.print()">Print</button></div>
<main id="dashboard"></main>
<script>
const MODEL={encoded};
const settings=MODEL.settings, root=document.getElementById('dashboard');
document.getElementById('dashboard-title').textContent=settings.title||'Chart Dashboard';
document.getElementById('dashboard-subtitle').textContent=settings.subtitle||'';
root.style.setProperty('--cols',settings.columns||2);
const palette=['#2457d6','#e45756','#2a9d8f','#f2a541','#7b61a8','#17a2b8','#8a9a5b'];
function sortRows(rows, mode){{const r=[...rows]; if(mode==='input')return r; const key=mode.startsWith('x')?'x':'y', dir=mode.endsWith('descending')?-1:1; return r.sort((a,b)=>dir*(key==='y'?a.y-b.y:String(a.x).localeCompare(String(b.x),undefined,{{numeric:true}})));}}
function groups(rows){{const m=new Map();for(const r of rows){{if(!m.has(r.series))m.set(r.series,[]);m.get(r.series).push(r)}}return m}}
function traceColor(rows,index){{return rows.find(r=>r.color)?.color||palette[index%palette.length]}}
function traces(chart, rows){{rows=sortRows(rows,chart.sort);if(['pie','donut'].includes(chart.chart_type)){{return [{{type:'pie',labels:rows.map(r=>r.x),values:rows.map(r=>r.y),text:rows.map(r=>r.label),hole:chart.chart_type==='donut'?chart.donut_hole:0,marker:{{colors:rows.map((r,i)=>r.color||palette[i%palette.length])}},hovertemplate:'%{{label}}<br>%{{value:,}} (%{{percent}})<extra></extra>'}}]}}
let out=[],i=0;for(const [name,rs] of groups(rows)){{const color=traceColor(rs,i++), base={{name,x:rs.map(r=>r.x),y:rs.map(r=>r.y),text:rs.map(r=>r.label),customdata:rs.map(r=>r.row),hovertemplate:'%{{x}}<br>%{{y:,}}<br>%{{text}}<extra>'+name+'</extra>'}};
if(chart.chart_type==='bar')out.push({{...base,type:'bar',marker:{{color}}}});
if(chart.chart_type==='line')out.push({{...base,type:'scatter',mode:'lines+markers',line:{{color,width:3}},marker:{{color}}}});
if(chart.chart_type==='area')out.push({{...base,type:'scatter',mode:'lines',stackgroup:chart.stacked?'one':undefined,fill:chart.stacked?undefined:'tozeroy',line:{{color,width:2}}}});
if(chart.chart_type==='scatter')out.push({{...base,type:'scatter',mode:'markers',marker:{{color,size:rs.map(r=>r.size||11),sizemode:'diameter',opacity:.8,line:{{color:'#fff',width:1}}}}}});}}return out}}
function table(rows){{if(!settings.show_data_table)return '';const body=rows.map(r=>`<tr><td>${{esc(r.series)}}</td><td>${{esc(r.x)}}</td><td>${{r.y}}</td><td>${{esc(r.label)}}</td></tr>`).join('');return `<details><summary>View source data (${{rows.length}} rows)</summary><table><thead><tr><th>Series</th><th>X</th><th>Y</th><th>Label</th></tr></thead><tbody>${{body}}</tbody></table></details>`}}
function esc(v){{const d=document.createElement('div');d.textContent=v??'';return d.innerHTML}}
function render(){{root.innerHTML='';for(const chart of MODEL.charts){{const rows=MODEL.data.filter(r=>r.chart_id===chart.chart_id), card=document.createElement('section');card.className='card';card.dataset.search=(chart.title+' '+chart.subtitle+' '+chart.chart_type).toLowerCase();card.innerHTML=`<div class="card-head"><h2>${{esc(chart.title)}}</h2><p class="subtitle">${{esc(chart.subtitle)}}</p></div><div class="plot" id="plot-${{CSS.escape(chart.chart_id)}}"></div>${{table(rows)}}`;root.append(card);const dark=document.documentElement.dataset.theme==='dark';const layout={{height:chart.height,margin:{{l:60,r:28,t:25,b:60}},paper_bgcolor:'transparent',plot_bgcolor:'transparent',font:{{color:dark?'#f3f4f6':'#172033'}},showlegend:chart.show_legend,legend:{{orientation:'h',y:-.18}},xaxis:{{title:chart.x_label,gridcolor:dark?'#374151':'#e7ecf3',automargin:true}},yaxis:{{title:chart.y_label,gridcolor:dark?'#374151':'#e7ecf3',automargin:true}},barmode:chart.stacked?'stack':'group'}};if(rows.length)Plotly.newPlot(card.querySelector('.plot'),traces(chart,rows),layout,{{responsive:settings.responsive,displaylogo:false,toImageButtonOptions:{{format:'png',filename:chart.chart_id,scale:2}}}});else card.querySelector('.plot').innerHTML='<div class="empty">No valid data</div>';}}}}
render();
document.getElementById('search').addEventListener('input',e=>{{const q=e.target.value.toLowerCase();document.querySelectorAll('.card').forEach(c=>c.hidden=!c.dataset.search.includes(q))}});
document.getElementById('theme').onclick=()=>{{document.documentElement.dataset.theme=document.documentElement.dataset.theme==='dark'?'light':'dark';render()}};
document.getElementById('download').onclick=()=>{{const blob=new Blob([JSON.stringify(MODEL,null,2)],{{type:'application/json'}}),a=document.createElement('a');a.href=URL.createObjectURL(blob);a.download='chart-dashboard-data.json';a.click();URL.revokeObjectURL(a.href)}};
</script></body></html>"""


def qa_markdown(source: Path, qa: QA, charts: list[dict[str, Any]], data: list[dict[str, Any]]) -> str:
    lines = [
        "# Chart Dashboard QA Report", "",
        f"- Source: `{source.name}`",
        f"- Generated: {datetime.now().astimezone().isoformat(timespec='seconds')}",
        f"- Valid charts: {len(charts)}",
        f"- Valid data rows: {len(data)}",
        f"- Errors: {len(qa.errors)}",
        f"- Warnings: {len(qa.warnings)}", "",
        "## Errors", "",
    ]
    lines += [f"- {e}" for e in qa.errors] or ["None."]
    lines += ["", "## Warnings", ""]
    lines += [f"- {w}" for w in qa.warnings] or ["None."]
    return "\n".join(lines) + "\n"


def render(source: Path, output: Path, strict: bool = False) -> QA:
    settings, charts, data, qa = parse_document(source)
    if qa.errors:
        raise InputError("\n".join(qa.errors))
    if strict and qa.warnings:
        raise InputError("Strict mode rejected warnings:\n" + "\n".join(qa.warnings))
    vendor = Path(__file__).parent / "vendor" / "plotly.min.js"
    if not vendor.exists():
        raise FileNotFoundError(f"Offline Plotly asset not found: {vendor}")
    output.mkdir(parents=True, exist_ok=True)
    bundle = payload(settings, charts, data)
    (output / "data.json").write_text(json.dumps(bundle, indent=2, ensure_ascii=False), encoding="utf-8")
    (output / "index.html").write_text(build_html(bundle, vendor.read_text(encoding="utf-8")), encoding="utf-8")
    (output / "qa_report.md").write_text(qa_markdown(source, qa, charts, data), encoding="utf-8")
    return qa


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="Word .docx containing SETTINGS, CHARTS and DATA tables")
    parser.add_argument("-o", "--output", type=Path, default=Path("output/chart_dashboard"))
    parser.add_argument("--strict", action="store_true", help="Fail if warnings are present")
    parser.add_argument("--version", action="version", version=VERSION)
    args = parser.parse_args()
    try:
        qa = render(args.input, args.output, args.strict)
    except (InputError, FileNotFoundError) as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 2
    print(f"Created {args.output / 'index.html'} ({len(qa.warnings)} warning(s))")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
